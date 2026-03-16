"""Markdown to DOCX conversion for marksmith.

Pipeline
--------
Markdown file  →  parse YAML front-matter  →  render Markdown body to HTML
→  walk HTML tree  →  write python-docx Document  →  save .docx

Supported Markdown elements
---------------------------
* Headings  H1 – H6
* Paragraphs with inline bold, italic, inline-code and links
* Fenced and indented code blocks
* Unordered and ordered lists (nested up to three levels)
* Block-quotes
* Tables (with bold header row)
* Thematic breaks (horizontal rules)

YAML Front-matter
-----------------
The Markdown file may start with a YAML front-matter block delimited by
``---`` lines.  Recognised keys are written to the DOCX core properties:

    ---
    title:          My Document
    version:        1.0
    author:         Paul Cummings
    date:           2026-03-16
    classification: Internal
    ---

All keys are also available as Jinja2 variables when using a template
(see ``marksmith.template``).

Template support
----------------
Pass ``--template company.docx`` to merge front-matter metadata into a branded
``.docx`` template and insert the Markdown body at a ``{{p marksmith_content }}``
placeholder.  Requires ``pip install marksmith[template]``.
See :mod:`marksmith.template` for full documentation.
"""

from __future__ import annotations

import contextlib
import os
from pathlib import Path

import frontmatter
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Markdown extensions enabled for all conversions.
_MD_EXTENSIONS = ["tables", "fenced_code", "sane_lists"]

# Maximum list-nesting depth supported by the default DOCX styles.
_MAX_LIST_DEPTH = 3

# Optional template support — available when marksmith[template] is installed.
try:
    from marksmith import template as _template_module
except ImportError:
    _template_module = None  # type: ignore[assignment]


# ── Public API ────────────────────────────────────────────────────────────────


def md_to_docx(
    input_path: str,
    output_path: str,
    template_path: str | None = None,
) -> None:
    """Convert a Markdown file to a DOCX document.

    Parameters
    ----------
    input_path:
        Path to the source ``.md`` file.
    output_path:
        Destination path for the generated ``.docx`` file.  The parent
        directory is created automatically if it does not exist.
    template_path:
        Optional path to a ``.docx`` template file containing Jinja2-style
        placeholders.  Requires ``pip install marksmith[template]``.
        Use ``{{p marksmith_content }}`` in the template to mark the content
        insertion point.  See :mod:`marksmith.template` for full details.
    """
    input_path = str(input_path)
    output_path = str(output_path)

    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    suffix = Path(output_path).suffix.lower()
    if suffix != ".docx":
        raise ValueError(f"Unsupported output format '{suffix}'. Only .docx is currently supported.")

    if template_path is not None:
        if _template_module is None:
            raise ImportError(
                "Template support requires the 'docxtpl' package. Install it with: pip install marksmith[template]"
            )
        _template_module.md_to_docx_templated(input_path, output_path, template_path)
        return

    with open(input_path, encoding="utf-8") as fh:
        raw = fh.read()

    metadata, body = _parse_frontmatter(raw)
    html = _md_to_html(body)
    doc = _html_to_docx(html)
    _apply_core_properties(doc, metadata)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ── Internal helpers ──────────────────────────────────────────────────────────


def _parse_frontmatter(raw: str) -> tuple[dict, str]:
    """Return ``(metadata_dict, markdown_body)`` parsed from *raw*."""
    post = frontmatter.loads(raw)
    return dict(post.metadata), post.content


def _md_to_html(md_text: str) -> str:
    """Render Markdown text to an HTML string."""
    return markdown.markdown(md_text, extensions=_MD_EXTENSIONS)


def _apply_core_properties(doc: Document, metadata: dict) -> None:
    """Write recognised front-matter keys to DOCX core properties."""
    props = doc.core_properties
    if "title" in metadata:
        props.title = str(metadata["title"])
    if "author" in metadata:
        props.author = str(metadata["author"])
    if "description" in metadata:
        props.description = str(metadata["description"])
    if "version" in metadata:
        # revision must be a positive integer; convert floats/strings best-effort.
        with contextlib.suppress(ValueError, TypeError):
            props.revision = int(float(str(metadata["version"])))
    if "keywords" in metadata:
        props.keywords = str(metadata["keywords"])


def _html_to_docx(html: str) -> Document:
    """Build and return a :class:`docx.Document` from an HTML string."""
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:
        _process_block(doc, element, depth=1)

    return doc


# ── Block-level element processors ───────────────────────────────────────────


def _process_block(doc: Document, element: Tag | NavigableString, depth: int = 1) -> None:
    """Dispatch a single block-level HTML element to the appropriate handler."""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            doc.add_paragraph(text)
        return

    tag = element.name

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        doc.add_heading(element.get_text(strip=True), level=level)

    elif tag == "p":
        para = doc.add_paragraph()
        _add_inline_content(para, element)

    elif tag in ("ul", "ol"):
        _process_list(doc, element, ordered=(tag == "ol"), depth=depth)

    elif tag == "pre":
        code_el = element.find("code")
        code_text = code_el.get_text() if code_el else element.get_text()
        para = doc.add_paragraph(style="Normal")
        run = para.add_run(code_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    elif tag == "blockquote":
        _process_blockquote(doc, element)

    elif tag == "table":
        _process_table(doc, element)

    elif tag == "hr":
        _add_horizontal_rule(doc)

    # Silently ignore unknown / structural tags (html, body, div, etc.)


def _process_list(
    doc: Document,
    element: Tag,
    ordered: bool = False,
    depth: int = 1,
) -> None:
    """Recursively render an HTML ``<ul>`` or ``<ol>`` into list paragraphs."""
    clamped = min(depth, _MAX_LIST_DEPTH)
    if ordered:
        style = "List Number" if clamped == 1 else f"List Number {clamped}"
    else:
        style = "List Bullet" if clamped == 1 else f"List Bullet {clamped}"

    for li in element.find_all("li", recursive=False):
        para = doc.add_paragraph(style=style)

        # Add inline content from the list item, skipping nested list children.
        for child in li.children:
            child_tag = getattr(child, "name", None)
            if child_tag not in ("ul", "ol"):
                _add_inline_content_node(para, child)

        # Recurse into any nested lists.
        for child in li.children:
            child_tag = getattr(child, "name", None)
            if child_tag in ("ul", "ol"):
                _process_list(doc, child, ordered=(child_tag == "ol"), depth=depth + 1)


def _process_blockquote(doc: Document, element: Tag) -> None:
    """Render a ``<blockquote>`` as indented italic paragraphs."""
    for child in element.children:
        child_tag = getattr(child, "name", None)
        if child_tag == "p":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            for run_child in child.children:
                run = _add_inline_content_node(para, run_child)
                if run:
                    run.italic = True
        elif isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.left_indent = Inches(0.4)
                for run in para.runs:
                    run.italic = True


def _process_table(doc: Document, element: Tag) -> None:
    """Render an HTML ``<table>`` into a DOCX table."""
    rows = element.find_all("tr")
    if not rows:
        return

    col_count = max(len(r.find_all(["td", "th"])) for r in rows)
    if col_count == 0:
        return

    table = doc.add_table(rows=len(rows), cols=col_count)
    with contextlib.suppress(KeyError):
        table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for col_idx, cell in enumerate(cells):
            docx_cell = table.rows[row_idx].cells[col_idx]
            docx_cell.text = ""
            para = docx_cell.paragraphs[0]
            _add_inline_content(para, cell)
            if cell.name == "th":
                for run in para.runs:
                    run.bold = True


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a paragraph styled as a horizontal rule using OOXML borders."""
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ── Inline content helpers ────────────────────────────────────────────────────


def _add_inline_content(para, element: Tag) -> None:
    """Add all inline children of *element* to *para*."""
    for child in element.children:
        _add_inline_content_node(para, child)


def _add_inline_content_node(para, node) -> None:  # noqa: PLR0911
    """Add a single inline node (NavigableString or Tag) to *para*.

    Returns the last Run added, or None if no run was created.
    """
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            return para.add_run(text)
        return None

    tag = node.name

    if tag in ("strong", "b"):
        run = para.add_run(node.get_text())
        run.bold = True
        return run

    if tag in ("em", "i"):
        run = para.add_run(node.get_text())
        run.italic = True
        return run

    if tag == "code":
        run = para.add_run(node.get_text())
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        return run

    if tag in ("del", "s"):
        run = para.add_run(node.get_text())
        run.font.strike = True
        return run

    if tag == "a":
        # TODO: Add proper OOXML hyperlink support (requires relationship injection).
        run = para.add_run(node.get_text())
        run.underline = True
        return run

    if tag == "img":
        # TODO: Embed images via run.add_picture() once image download/path
        # resolution is implemented.
        run = para.add_run(f"[image: {node.get('alt', '')}]")
        run.italic = True
        return run

    # For any other inline tag (span, abbr, etc.) recurse into its children.
    for child in node.children:
        _add_inline_content_node(para, child)
    return None
